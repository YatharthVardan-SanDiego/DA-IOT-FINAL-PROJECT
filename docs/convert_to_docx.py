"""
Convert docs/report.md to docs/report.docx (Google Docs / Word compatible).

Usage:
    pip install python-docx
    python docs/convert_to_docx.py

Output: docs/report.docx
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_MD = os.path.join(SCRIPT_DIR, "report.md")
OUTPUT_DOCX = os.path.join(SCRIPT_DIR, "report.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_page_margins(doc, inches=1.0):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def set_body_font(doc, name="Times New Roman", size=12):
    style = doc.styles["Normal"]
    font = style.font
    font.name = name
    font.size = Pt(size)


def add_horizontal_rule(doc):
    """Add a thin bottom-border paragraph to simulate a horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def apply_inline(run_adder, text):
    """
    Parse inline markdown (bold, italic, code, bold+italic) and add runs
    via run_adder(text, bold, italic, code).
    """
    # Pattern order matters: bold+italic first, then bold, then italic, then code
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"   # bold+italic
        r"|(\*\*(.+?)\*\*)"       # bold
        r"|(\*(.+?)\*)"           # italic
        r"|(`(.+?)`)"             # inline code
    )
    last = 0
    for m in pattern.finditer(text):
        # Plain text before this match
        if m.start() > last:
            run_adder(text[last:m.start()], False, False, False)
        if m.group(1):   # bold+italic
            run_adder(m.group(2), True, True, False)
        elif m.group(3): # bold
            run_adder(m.group(4), True, False, False)
        elif m.group(5): # italic
            run_adder(m.group(6), False, True, False)
        elif m.group(7): # code
            run_adder(m.group(8), False, False, True)
        last = m.end()
    if last < len(text):
        run_adder(text[last:], False, False, False)


def add_inline_paragraph(doc, text, style="Normal", alignment=None):
    """Add a paragraph with inline markdown formatting applied."""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment

    def add_run(t, bold, italic, code):
        run = p.add_run(t)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Courier New"
            run.font.size = Pt(10)

    apply_inline(add_run, text)
    return p


def add_table_from_rows(doc, rows):
    """
    rows: list of lists of strings (already split on |).
    First row = header, second row = separator (ignored), rest = data.
    """
    if len(rows) < 2:
        return

    # Determine if row[1] is a separator (contains only -, :, |, space)
    data_rows = rows[:]
    header_row = data_rows[0]
    if len(data_rows) > 1 and all(
        re.match(r"^[-: ]+$", cell) for cell in data_rows[1]
    ):
        data_rows = [data_rows[0]] + data_rows[2:]

    col_count = max(len(r) for r in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= col_count:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            is_header = r_idx == 0

            def add_run(t, bold, italic, code, _is_header=is_header):
                run = p.add_run(t)
                run.bold = bold or _is_header
                run.italic = italic
                if code:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)

            apply_inline(add_run, cell_text.strip())

    # Pad any short rows
    for r_idx, row in enumerate(data_rows):
        for c_idx in range(len(row), col_count):
            table.cell(r_idx, c_idx).text = ""

    doc.add_paragraph()  # spacing after table


def parse_table_row(line):
    """Split a markdown table row into cells, stripping leading/trailing |."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c for c in line.split("|")]


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def convert(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    set_page_margins(doc)
    set_body_font(doc)

    # Letter page size is the default for python-docx; keep it.

    i = 0
    table_buffer = []   # accumulate table rows

    def flush_table():
        if table_buffer:
            add_table_from_rows(doc, table_buffer)
            table_buffer.clear()

    while i < len(lines):
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()

        # ── Table row ──────────────────────────────────────────────────────
        if stripped.startswith("|"):
            table_buffer.append(parse_table_row(stripped))
            i += 1
            continue
        else:
            flush_table()

        # ── Blank line ─────────────────────────────────────────────────────
        if stripped == "":
            i += 1
            continue

        # ── Horizontal rule ────────────────────────────────────────────────
        if re.match(r"^---+$", stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── ATX Headings ───────────────────────────────────────────────────
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            # Strip markdown bold/italic from heading text for simplicity
            heading_text = re.sub(r"\*+(.+?)\*+", r"\1", heading_text)
            heading_text = re.sub(r"`(.+?)`", r"\1", heading_text)
            doc.add_heading(heading_text, level=level)
            i += 1
            continue

        # ── Bullet list ────────────────────────────────────────────────────
        bullet_match = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet_match:
            add_inline_paragraph(doc, bullet_match.group(1), style="List Bullet")
            i += 1
            continue

        # ── Numbered list ──────────────────────────────────────────────────
        num_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if num_match:
            add_inline_paragraph(doc, num_match.group(1), style="List Number")
            i += 1
            continue

        # ── Image (skip, add caption placeholder) ─────────────────────────
        if re.match(r"^!\[", stripped):
            alt_match = re.match(r"^!\[([^\]]*)\]", stripped)
            alt = alt_match.group(1) if alt_match else "Figure"
            p = doc.add_paragraph(f"[Figure: {alt}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            i += 1
            continue

        # ── Italicised block-level lines (e.g. *Table N.* captions, *Note.*) ─
        # These are normal paragraphs — handled by add_inline_paragraph below.

        # ── Fenced code block ──────────────────────────────────────────────
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            # Light grey shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)
            continue

        # ── Regular paragraph ──────────────────────────────────────────────
        add_inline_paragraph(doc, stripped)
        i += 1

    flush_table()

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    convert(INPUT_MD, OUTPUT_DOCX)
